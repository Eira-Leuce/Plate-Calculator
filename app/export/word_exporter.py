# app/export/word_exporter.py
from docx import Document
from docx.shared import Inches
import io


def export_to_word(filepath, figures_dict, text_results, main_gui_figure=None):
    """
    Экспортирует графики и текстовые результаты в Word документ.
    figures_dict: словарь, где ключ - заголовок, значение - объект matplotlib Figure.
    text_results: строка с текстовыми результатами.
    main_gui_figure: основная фигура из GUI, чтобы попытаться взять оттуда стиль (не реализовано глубоко).
    """
    doc = Document()
    doc.add_heading('Результаты расчета изгиба плиты', level=1)

    if text_results:
        doc.add_heading('Числовые результаты', level=2)
        doc.add_paragraph(text_results)
        doc.add_page_break()

    if figures_dict:
        doc.add_heading('Графики', level=2)
        for title, fig_object in figures_dict.items():
            if fig_object:  # Убедимся, что фигура существует
                doc.add_heading(title, level=3)

                # Сохраняем фигуру во временный байтовый поток
                img_stream = io.BytesIO()
                try:
                    fig_object.savefig(img_stream, format='png', dpi=150)  # Уменьшим dpi для файла
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(6.0))
                except Exception as e:
                    print(f"Ошибка при сохранении графика '{title}' в Word: {e}")
                    doc.add_paragraph(f"(Ошибка при генерации графика: {title})")
                finally:
                    img_stream.close()

    doc.save(filepath)
